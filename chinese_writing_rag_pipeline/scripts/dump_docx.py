#!/usr/bin/env python3
"""Dump full structure of a docx: paragraphs (with style), tables, and comments."""
import sys
import zipfile
import re
from docx import Document
from docx.oxml.ns import qn

def get_comments(path):
    """Extract comments.xml content if present, mapped by comment id."""
    comments = {}
    try:
        with zipfile.ZipFile(path) as z:
            names = z.namelist()
            if 'word/comments.xml' in names:
                xml = z.read('word/comments.xml').decode('utf-8', errors='replace')
                # crude extraction of comment id + text
                for m in re.finditer(r'<w:comment[^>]*w:id="(\d+)"[^>]*>(.*?)</w:comment>', xml, re.S):
                    cid, body = m.groups()
                    texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', body, re.S)
                    comments[cid] = ''.join(texts)
    except Exception as e:
        print(f"  [comment extraction error: {e}]", file=sys.stderr)
    return comments

def dump(path):
    print(f"\n{'='*100}\nFILE: {path}\n{'='*100}")
    doc = Document(path)
    comments = get_comments(path)
    if comments:
        print(f"\n--- COMMENTS FOUND: {len(comments)} ---")
        for cid, text in comments.items():
            print(f"  [comment {cid}]: {text}")

    print(f"\n--- BODY (paragraphs + tables in document order) ---")
    body = doc.element.body
    para_idx = 0
    table_idx = 0
    paras = doc.paragraphs
    tables = doc.tables

    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = paras[para_idx]
            para_idx += 1
            text = p.text
            style = p.style.name if p.style else ""
            # check for comment references
            comment_refs = child.findall('.//' + qn('w:commentReference'))
            cref_ids = [c.get(qn('w:id')) for c in comment_refs]
            marker = f" [COMMENT_REF:{cref_ids}]" if cref_ids else ""
            if text.strip() or marker:
                print(f"P[{para_idx-1}] ({style}): {text}{marker}")
        elif child.tag == qn('w:tbl'):
            if table_idx < len(tables):
                t = tables[table_idx]
                print(f"\n  TABLE[{table_idx}] ({len(t.rows)} rows x {len(t.columns)} cols):")
                for ri, row in enumerate(t.rows):
                    cells = [c.text.replace('\n', ' | ') for c in row.cells]
                    print(f"    row{ri}: {cells}")
                table_idx += 1
            print()

if __name__ == '__main__':
    for p in sys.argv[1:]:
        dump(p)
